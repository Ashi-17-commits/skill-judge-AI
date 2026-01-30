from docx import Document
from pathlib import Path

p = Path(__file__).parent / "test_resume.docx"

doc = Document()
doc.add_heading('John Doe', level=1)
doc.add_paragraph('Software Engineer with experience in Python and SQL.')
doc.add_heading('Skills', level=2)
doc.add_paragraph('Python, SQL, Docker, AWS')
doc.add_heading('Experience', level=2)
doc.add_paragraph('Worked on backend services, improved performance by 30%.')

doc.save(p)
print('Wrote', p)
