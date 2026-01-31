from dataclasses import dataclass
from pathlib import Path
from typing import Literal

import pdfplumber
from docx import Document


SupportedExtension = Literal[".pdf", ".docx"]


@dataclass
class ExtractResult:
    """Result of resume text extraction including metadata for scoring."""

    text: str
    page_count: int


def detect_extension(file_path: Path) -> SupportedExtension:
    """
    Validate and normalize a file extension for supported resume types.
    """

    ext = file_path.suffix.lower()
    if ext not in {".pdf", ".docx"}:
        raise ValueError("Unsupported file type. Only PDF and DOCX are allowed.")
    return ext  # type: ignore[return-value]


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract plain text from a PDF using pdfplumber.

    The function concatenates text from all pages and strips trailing
    whitespace while preserving basic line breaks for downstream parsing.
    """

    text_parts: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract plain text from a DOCX file using python-docx.
    """

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


def extract_text(file_path: Path) -> str:
    """
    Dispatch to the correct extraction routine based on file extension.
    """

    return extract_with_meta(file_path).text


def extract_with_meta(file_path: Path) -> ExtractResult:
    """
    Extract plain text and page count (for PDF) or estimated pages (DOCX).
    Used by ATS scoring for format/structure signals.
    """

    ext = detect_extension(file_path)
    if ext == ".pdf":
        text_parts: list[str] = []
        page_count = 0
        with pdfplumber.open(str(file_path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        return ExtractResult(text=text, page_count=page_count)
    if ext == ".docx":
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs).strip()
        # Estimate pages: ~30 lines per page for typical resume
        line_count = len([p for p in doc.paragraphs if p.text.strip()])
        page_count = max(1, (line_count + 29) // 30)
        return ExtractResult(text=text, page_count=page_count)
    raise ValueError(f"Unsupported file extension: {ext}")
